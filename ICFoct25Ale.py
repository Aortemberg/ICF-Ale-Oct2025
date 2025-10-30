# ICFAle.py
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
import io
import zipfile
import re
from docx.oxml import OxmlElement

# -----------------------------
# Configuración de la aplicación Streamlit
# -----------------------------
st.set_page_config(page_title="Generador DOCX Consentimientos", layout="wide")

st.title("🩺 Generador automático de Consentimientos (Excel → Word)")

st.markdown("""
Subí tu **modelo.docx** (plantilla con placeholders `<<...>>`) y el **datos.xlsx** con la información de cada investigador.  
El nombre del archivo final se construirá con el Investigador, el Nro. de Centro y el número de protocolo.
""")

# Cargadores de archivos
uploaded_docx = st.file_uploader("📄 Subí el documento modelo (.docx)", type=["docx"])
uploaded_xlsx = st.file_uploader("📊 Subí el Excel (.xlsx)", type=["xlsx"])

# Variables globales para lógica de reemplazo
texto_anticonceptivo_original = (
    "El médico del estudio discutirá con usted qué método anticonceptivo se considera adecuado. "
    "El patrocinador y/o el investigador del estudio garantizarán su acceso al método anticonceptivo "
    "acordado y necesario para su participación en este estudio"
)

texto_ba_reemplazo = (
    "El médico del estudio discutirá con usted qué métodos anticonceptivos se consideran adecuados. "
    "El Patrocinador y/o el médico del estudio garantizará su acceso a este método anticonceptivo "
    "acordado y necesario para su participación en el ensayo. El costo de los métodos anticonceptivos "
    "seleccionados correrá a cargo del Patrocinador."
)

# -----------------------------
# Funciones auxiliares
# -----------------------------
def remove_paragraph(paragraph):
    """Elimina un párrafo completamente del documento."""
    p = paragraph._element
    p.getparent().remove(p)

def replace_text_in_runs(paragraph, old, new):
    """Reemplaza texto en fragmentos de párrafo (runs) sin romper el formato original."""
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)

def replace_text_in_doc(doc, replacements):
    """Aplica reemplazos en todos los párrafos y tablas del documento."""
    def process_paragraphs(paragraphs):
        for p in paragraphs:
            for old, new in replacements.items():
                replace_text_in_runs(p, old, new)
            fulltext = p.text
            for old, new in replacements.items():
                if old in fulltext:
                    for r in p.runs:
                        r.text = ""
                    p.add_run(fulltext.replace(old, new))
    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

def find_paragraphs_containing(doc, snippet):
    """Busca y devuelve todos los párrafos que contienen el fragmento de texto dado."""
    res = []
    for p in doc.paragraphs:
        if snippet.lower() in p.text.lower():
            res.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if snippet.lower() in p.text.lower():
                        res.append(p)
    return res

def get_docx_creation_date(file):
    """Lee la fecha de modificación del modelo Word (si existe) o usa la fecha actual."""
    try:
        from zipfile import ZipFile
        from xml.etree import ElementTree as ET
        file.seek(0)
        with ZipFile(file) as docx:
            core = docx.read("docProps/core.xml")
            tree = ET.fromstring(core)
            ns = {"dcterms": "http://purl.org/dc/terms/"}
            modified = tree.find("dcterms:modified", ns)
            if modified is not None and modified.text:
                dt = datetime.fromisoformat(modified.text.replace("Z", "+00:00"))
                return dt.strftime("%d/%m/%Y")
    except Exception:
        pass
    return datetime.now().strftime("%d/%m/%Y")

def set_global_font_style(doc, font_name="Arial", font_size=11, font_color=RGBColor(0, 0, 0)):
    """Aplica formato de fuente consistente a todo el documento, incluyendo tablas."""
    font_size_pt = Pt(font_size)
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = font_name
            run.font.size = font_size_pt
            run.font.color.rgb = font_color
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = font_name
                        run.font.size = font_size_pt
                        run.font.color.rgb = font_color

# -----------------------------
# Procesamiento de cada fila
# -----------------------------
def process_row_and_generate_doc(template_bytes, row, fecha_modelo):
    doc = Document(io.BytesIO(template_bytes))

    replacements = {
        "<<NUMERO_PROTOCOLO>>": str(row.get("Numero de protocolo", "")).strip(),
        "<<TITULO_ESTUDIO>>": str(row.get("Titulo del Estudio", "")).strip(),
        "<<PATROCINADOR>>": str(row.get("Patrocinador", "")).strip(),
        "<<INVESTIGADOR>>": str(row.get("Investigador", "")).strip(),
        "<<INSTITUCION>>": str(row.get("Institucion", "")).strip(),
        "<<DIRECCION>>": str(row.get("Direccion", "")).strip(),
        "<<CARGO_INVESTIGADOR>>": str(row.get("Cargo del Investigador en la Institucion", "")).strip(),
        "<<Centro_Nro.>>": str(row.get("Nro. de Centro", "")).strip(),
        "<<COMITE>>": str(row.get("COMITE", "")).strip(),
        "<<SUBINVESTIGADOR>>": str(row.get("Subinvestigador", "")).strip(),
        "<<TELEFONO_24HS>>": str(row.get("TELEFONO 24HS", "")).strip(),
        "<<TELEFONO_24HS_SUBINV>>": str(row.get("TELEFONO 24HS subinvestigador", "")).strip(),
    }

    # ✅ Si Subinvestigador está vacío → eliminar párrafos y placeholders relacionados
    sub_val = replacements.get("<<SUBINVESTIGADOR>>", "")
    if not sub_val:
        for key in ["<<SUBINVESTIGADOR>>", "<<TELEFONO_24HS_SUBINV>>"]:
            replacements.pop(key, None)
            for p in find_paragraphs_containing(doc, key):
                remove_paragraph(p)

    # Reemplazar placeholders
    replace_text_in_doc(doc, replacements)

    # Aplicar formato Arial 11 negro a todo
    set_global_font_style(doc)

    # Agregar línea final con fecha de modelo
    doc.add_paragraph(f"Documento basado en modelo de fecha: {fecha_modelo}")

    out_io = io.BytesIO()
    doc.save(out_io)
    out_io.seek(0)
    return out_io

# -----------------------------
# Ejecución principal
# -----------------------------
if uploaded_docx and uploaded_xlsx:
    uploaded_docx.seek(0)
    fecha_modelo = get_docx_creation_date(uploaded_docx)

    try:
        df = pd.read_excel(uploaded_xlsx, engine="openpyxl")
        if df.empty:
            st.error("⚠️ El archivo Excel está vacío.")
            st.stop()
    except Exception as e:
        st.error(f"Error leyendo el Excel: {e}")
        st.stop()

    uploaded_docx.seek(0)
    template_bytes = uploaded_docx.read()
    zip_io = io.BytesIO()

    with st.spinner("⏳ Generando documentos..."):
        with zipfile.ZipFile(zip_io, "w", zipfile.ZIP_DEFLATED) as zf:
            for idx, row in df.iterrows():
                try:
                    doc_io = process_row_and_generate_doc(template_bytes, row.to_dict(), fecha_modelo)
                except Exception as e:
                    st.error(f"Error procesando fila {idx + 2}: {e}")
                    continue

                inv = str(row.get("Investigador", "")).strip()
                centro = str(row.get("Nro. de Centro", "")).strip()
                protocolo = str(row.get("Numero de protocolo", "")).strip()

                safe_inv = re.sub(r'[\\/*?:"<>|]', "_", inv)[:100]
                safe_centro = re.sub(r'[\\/*?:"<>|]', "_", centro)[:50]
                safe_prot = re.sub(r'[\\/*?:"<>|]', "_", protocolo)[:50]

                filename = f"{safe_inv} - Centro {safe_centro} - {safe_prot}.docx"
                if not safe_inv and not safe_centro:
                    filename = f"documento_generado_{idx + 1}.docx"

                zf.writestr(filename, doc_io.getvalue())

    zip_io.seek(0)
    st.success(f"✅ ¡Documentos generados correctamente! ({len(df)} archivos)")
    st.download_button(
        "📥 Descargar ZIP",
        data=zip_io.getvalue(),
        file_name="consentimientos_generados.zip",
        mime="application/zip"
    )
else:
    st.info("👆 Subí el modelo .docx y el .xlsx para comenzar.")
